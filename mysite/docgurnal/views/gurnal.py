from io import BytesIO
import os
import re
from typing import Any
from django.contrib.auth.decorators import login_required
from django.db.models.query import QuerySet
from django.http import HttpResponse
from django.urls import reverse_lazy
from django.utils.decorators import method_decorator
from django.views.generic import ListView, UpdateView, DetailView, CreateView, DeleteView
import docx
from mysite.docgurnal.models import *
item_for_page = 15

def createDoc(request, doc_id):

    doc_with_id = gurnal.objects.get(pk=doc_id)
    doc_fields = dict(doc_with_id.get_fields())

    user = User.objects.get(pk=doc_with_id.user.id)

    path_to_doc = os.path.join(f'{os.getcwd()}/media/Zhurnal_praktiki.docx')
    Doc = docx.Document(path_to_doc)
    doc_tables = Doc.tables
    table = doc_tables[0]
    doc_slov = doc_fields

    for paragraph in Doc.paragraphs:
        for k, v in doc_slov.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)
    for table in Doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in doc_slov.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)

    cleaned_time_string = re.sub(r'[^\w\-_. ]', '_', str(doc_id))

    path_to_save = f'{os.getcwd()}/media/Zhurnal_praktiki-{user.username}-{cleaned_time_string}.docx'
    file_name = f'Zhurnal_praktiki-{user.username}-{cleaned_time_string}.docx'
    Doc.save(path_to_save)

    doc_data = BytesIO()
    Doc.save(doc_data)
    doc_data.seek(0)


def downloadDoc(request):
    doc_id = request.GET['id']
    doc_with_id = gurnal.objects.get(pk=doc_id)
    doc_fields = dict(doc_with_id.get_fields())
    user = User.objects.get(pk=doc_with_id.user.id)
    cleaned_time_string = re.sub(r'[^\w\-_. ]', '_', doc_id)
    path_to_save = f'{os.getcwd()}/media/Zhurnal_praktiki-{user.username}-{cleaned_time_string}.docx'
    file_name = f'Zhurnal_praktiki-{user.username}-{cleaned_time_string}.docx'
    with open(path_to_save, 'rb') as file:    
        response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'    
        return response

@method_decorator(login_required(login_url='/login/'), name='dispatch')
class gurnalView(ListView):
    model = gurnal
    context_object_name = 'gurnal_list'
    success_url = reverse_lazy('docgurnal: gurnal')

    paginate_by = item_for_page
    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(gurnalView,self).get_context_data(**kwargs)
        context['sometry'] = self.model._meta.verbose_name_plural
        context['col1name'] = self.model._meta.get_field("fiostud").verbose_name
        context['col2name'] = self.model._meta.get_field("facul").verbose_name
        context['col3name'] = self.model._meta.get_field("kafedra").verbose_name
        context['col4name'] = self.model._meta.get_field("groupnumber").verbose_name
        context['col5name'] = self.model._meta.get_field("napravlenie").verbose_name
        context['col6name'] = self.model._meta.get_field("vidpract").verbose_name
        context['col7name'] = self.model._meta.get_field("rukpractic").verbose_name
        context['col8name'] = self.model._meta.get_field("familinic").verbose_name
        context['col9name'] = self.model._meta.get_field("datasdacha").verbose_name
        context['col10name'] = self.model._meta.get_field("datanach").verbose_name
        context['col11name'] = self.model._meta.get_field("datakonec").verbose_name
        context['col12name'] = self.model._meta.get_field("predpriat").verbose_name
        context['col13name'] = self.model._meta.get_field("zadanie").verbose_name
        context['col14name'] = self.model._meta.get_field("dataotziv").verbose_name

        context['collastname'] = 'Сервисы'
        return context
    
    def get_queryset(self):
        return self.model.objects.filter(user=self.request.user)


@method_decorator(login_required(login_url='/login/'), name='dispatch')
class gurnalUpdate(UpdateView):
    model = gurnal
    template_name_suffix = '_update_form'
    fields = '__all__'
    success_url = '/docgurnal/gurnal/'

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(gurnalUpdate,self).get_context_data(**kwargs)
        context['sometry'] = self.model._meta.verbose_name
        context['slovar'] = {'fiostud','facul','kafedra','groupnumber','napravlenie','vidpract','rukpractic'}
        context['secslovar'] = {'familinic','datasdacha','datanach','datakonec','predpriat','zadanie','dataotziv'}
        context['hidenslovar'] = {'id'}
        context['dopslovar'] = {'created_at','edesc','erem','updated_at','created_it','updated_it'}
        context['model'] = self.model
        return context
    
    def form_valid(self, form):
        user = self.request.user
        form.instance.user = user
        # Сохраняем объект
        self.object = form.save()
        # Получаем id сохранённого объекта
        doc_id = self.object.id
        # Вызываем функцию createDoc с передачей id документа
        createDoc(self.request, doc_id)
        return super().form_valid(form)

@method_decorator(login_required(login_url='/login/'), name='dispatch')
class gurnalDetail(DetailView):
    model = gurnal
    context_object_name = 'gurnal_one'
    success_url = '/docgurnal/gurnal/'

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(gurnalDetail,self).get_context_data(**kwargs)
        context['sometry'] = self.model._meta.verbose_name
        context['slovar'] = {'fiostud','facul','kafedra','groupnumber','napravlenie','vidpract','rukpractic'}
        context['secslovar'] = {'familinic','datasdacha','datanach','datakonec','predpriat','zadanie','dataotziv'}
        context['hidenslovar'] = {'id'}
        context['dopslovar'] = {'created_at','edesc','erem','updated_at','created_it','updated_it'}
        context['model'] = self.model
        return context

@method_decorator(login_required(login_url='/login/'), name='dispatch')
class gurnalCreate(CreateView):
    model = gurnal
    context_object_name = 'gurnal_one'
    success_url = '/docgurnal/gurnal/'

    template_name_suffix = '_create_form'
    fields = '__all__'
    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(gurnalCreate,self).get_context_data(**kwargs)
        context['sometry'] = self.model._meta.verbose_name
        context['model'] = self.model
        context['slovar'] = {'fiostud','facul','kafedra','groupnumber','napravlenie','vidpract','rukpractic'}
        context['secslovar'] = {'familinic','datasdacha','datanach','datakonec','predpriat','zadanie','dataotziv'}
        context['hidenslovar'] = {'id'}
        context['dopslovar'] = {'created_at','edesc','erem','updated_at','created_it','updated_it'}
        return context
    
    def form_valid(self, form):
        user = self.request.user
        form.instance.user = user
        # Сохраняем объект
        self.object = form.save()
        # Получаем id сохранённого объекта
        doc_id = self.object.id
        # Вызываем функцию createDoc с передачей id документа
        createDoc(self.request, doc_id)
        return super().form_valid(form)

@method_decorator(login_required(login_url='/login/'), name='dispatch')
class gurnalDelete(DeleteView):
    model = gurnal
    success_url = '/docgurnal/gurnal/'


